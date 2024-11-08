from io import BytesIO                      
from docx import Document                   
from docx.shared import Pt, Inches, RGBColor  
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from PIL import Image                        
import base64                                
from openpyxl import load_workbook           

def create_document(images_base64, base64_img_first, file, word_doc_path):
    """Create a DOCX document with a provided base64 image first and extracted base64 images."""
    doc = Document()
    image_data = base64.b64decode(base64_img_first)
    image_stream = BytesIO(image_data)
        # Open image with PIL to determine size
    with Image.open(image_stream) as img:
        max_width = 6.0  # Max width in inches
        width, height = img.size
        aspect_ratio = width / height

        if width > height:
            adjusted_width = min(max_width, width / 96)  # Convert pixels to inches (assuming 96 dpi)
            adjusted_height = adjusted_width / aspect_ratio
        else:
            adjusted_height = min(max_width, height / 96)
            adjusted_width = adjusted_height * aspect_ratio

        # Add the image to the document
        doc.add_picture(image_stream, width=Inches(adjusted_width), height=Inches(adjusted_height))


    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')


    file_extension_school = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_school == 'xlsx':
        workbook_school = load_workbook(file)
        cover_page_school = workbook_school['Cover Page']
        diploma_text = cover_page_school['C6'].value.upper()
        if diploma_text.startswith("IT"):
            it_paragraph = doc.add_paragraph("Diploma in Information Technology")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("BM"):
            it_paragraph = doc.add_paragraph("Diploma in Business Management")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("DM"):
            it_paragraph = doc.add_paragraph("Diploma in Design & Media")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("HS"):
            it_paragraph = doc.add_paragraph("Diploma in Health & Social Sciences")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("EG"):
            it_paragraph = doc.add_paragraph("Diploma in Engineering")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        else:
            it_paragraph = doc.add_paragraph("Diploma in NYP")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'


    course = doc.add_paragraph('Course Syllabi for')
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = course.runs[0]
    run1.font.size = Pt(24)
    run1.font.name = 'Arial'


    file_extension = file.filename.rsplit('.', 1)[1].lower()
    if file_extension == 'xlsx':
        workbook = load_workbook(file)
        cover_page = workbook['Cover Page']
        diploma_text = cover_page['C7'].value.upper()

    diploma = doc.add_paragraph(diploma_text)
    diploma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = diploma.runs[0]
    run2.font.size = Pt(24)
    run2.font.name = 'Arial'
    run2.bold = True


    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    file_extension_year = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_year == 'xlsx':
        workbook_year = load_workbook(file)
        cover_page_year = workbook_year['Cover Page']
        diploma_text_year = cover_page_year['C12'].value.upper()
        year = doc.add_paragraph(diploma_text_year)
        year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = diploma.runs[0]
        run3.font.size = Pt(18)
        run3.font.name = 'Arial'


    doc.add_paragraph('')
    doc.add_paragraph('')


    revised = doc.add_paragraph('Date Revised: Apr  2023 v1.0')
    run4 = revised.runs[0]
    run4.font.size = Pt(11)
    run4.font.name = 'Arial'


    title1 = doc.add_paragraph(diploma_text)
    run5 = title1.runs[0]
    run5.font.size = Pt(16)
    run5.font.name = 'Arial'
    run5.bold = True
        

    title2 = doc.add_paragraph('Course Aims')
    run8 = title2.runs[0]
    run8.font.size = Pt(16)
    run8.font.name = 'Arial'
    run8.bold = True

    start_section = "Course Aims"
    end_section = "Course Learning Outcomes"
    text_between_sections = extract_text_between_sections(word_doc_path, start_section, end_section)
    
    course_paragraph = doc.add_paragraph(text_between_sections)
    course_run = course_paragraph.runs[0]
    course_run.font.size = Pt(12)
    course_run.font.name = 'Arial'


    title3 = doc.add_paragraph('Course Learning Outcomes')
    run10 = title3.runs[0]
    run10.font.size = Pt(16)
    run10.font.name = 'Arial'
    run10.bold = True

    course_paragraph_outcome = doc.add_paragraph("The competencies of a graduate are synthesized into 9 Course Competencies (CCs) as listed below.")
    course_run_outcome = course_paragraph_outcome.runs[0]
    course_run_outcome.font.size = Pt(12)
    course_run_outcome.font.name = 'Arial'


    for image_filename, image_b64 in images_base64:

        title2 = doc.add_paragraph('Course Competency Map')
        run8 = title2.runs[0]
        run8.font.size = Pt(16)
        run8.font.name = 'Arial'
        run8.bold = True
        # Decode base64 image and open with PIL
        image_data = base64.b64decode(image_b64)
        image_stream = BytesIO(image_data)

        # Open image with PIL to determine size
        with Image.open(image_stream) as img:
            max_width = 6.0  # Max width in inches
            width, height = img.size
            aspect_ratio = width / height

            if width > height:
                adjusted_width = min(max_width, width / 96)  # Convert pixels to inches (assuming 96 dpi)
                adjusted_height = adjusted_width / aspect_ratio
            else:
                adjusted_height = min(max_width, height / 96)
                adjusted_width = adjusted_height * aspect_ratio

            doc.add_picture(image_stream, width=Inches(adjusted_width), height=Inches(adjusted_height))



    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Graduate Attribute (GA)'
    hdr_cells[1].text = 'Course Learning Outcomes (CLOs)'
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Professional Proficiency'
    row_cells[1].text = 'CLO1: Apply technical knowledge and programming skills in the capacity of a business analytics IT professional.\n\nCLO2: Apply artificial intelligence and analytics technologies and tools to integrate technical and business knowledge to provide solution.\n\nCLO3: Demonstrate competence in artificial intelligence and analytics and be able to integrate and apply it effectively in different industry & domain.'

# Row 2: Competent in 21st Century Skills
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Competent in 21st Century Skills'
    row_cells[1].text = 'CLO4: Display the abilities to stay relevant by demonstrating independent learning, self-awareness and mental resilience, and personal effectiveness.\n\nCLO5: Demonstrate interpersonal skills and global perspectives by communicating and working effectively with people from diverse backgrounds.'

# Row 3: Innovative and Enterprising
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Innovative and Enterprising'
    row_cells[1].text = 'CLO6: Apply innovative and enterprising practices to achieve intended goals and drive continuous improvement with an interdisciplinary approach.'

# Row 4: Socially Responsible
    row_cells = table.rows[4].cells
    row_cells[0].text = 'Socially Responsible'
    row_cells[1].text = 'CLO7: Display personal and professional values and ethics by demonstrating inclusivity and responsibility towards the community, nation, and the world, and considering impact of actions and decisions on sustainability.'
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.runs
                for r in run:
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)
    doc.add_page_break()


    title_competency = doc.add_paragraph('Competency Canvases')
    run_course = title_competency.runs[0]
    run_course.font.size = Pt(16)
    run_course.font.name = 'Arial'
    run_course.bold = True
    doc.add_paragraph('')



    para5 = doc.add_paragraph('The learning outcomes of the diploma are to educate and train students who, by the time of successful completion of course, will be able to:')
    run13 = para5.runs[0]
    run13.font.size = Pt(12)
    run13.font.name = 'Arial'


    table = doc.add_table(rows=10, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'CC#'
    hdr_cells[1].text = 'Course Competencies (CCs)'

    # Manually filling in the content for each row
    content = [
        ('CC1', 'Data Visualisation & Journalism\n\nLearners will be competent in applying data visualisation and journalism techniques and tools to communicate data insights effectively with stakeholders to support business needs.'),
        ('CC2', 'Analytics & Computational Modelling\n\nLearners will be competent in applying data analytics and computational modelling skills using tools and algorithms to solve business problems.'),
        ('CC3', 'Applied Artificial Intelligence (AI)\n\nLearners will be competent in applying AI to build intelligent machine reasoning applications that derive hidden patterns and support decision-making.'),
        ('CC4', 'Data Administration & Management\n\nLearners will be competent in Big Data administration and management through data modelling and data manipulation techniques to meet business requirements.'),
        ('CC5', 'Analytics with Programming\n\nLearners will be competent in developing IT & data analytics applications according to users’ and business needs.'),
        ('CC6', 'Data Strategy & Design\n\nLearners will be competent to design robust data strategies to manage Big Data platforms aligned to stakeholders’ business values.'),
        ('CC7', 'Emerging Technology & Applications\n\nLearners will be competent to synthesise and integrate different emerging technology trends and developments to value-add and provide solutions for businesses.'),
        ('CC8', 'Business Needs Analysis & Strategy\n\nLearners will be competent to apply the business needs analysis and strategy skills to deliver service excellence for customers with diverse backgrounds.'),
        ('CC9', 'Data Security & Governance\n\nLearners will be well-versed in data security and governance and competent to apply cybersecurity principles to uphold personal and professional ethics.')
    ]

    # Fill each row with the CC content
    for i, (cc, desc) in enumerate(content, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = cc
        row_cells[1].text = desc

    # Optional: format font and style of the table content (for example, Arial, size 12)
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.runs
                for r in run:
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)

    # Align text vertically and horizontally (optional)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


    competence = doc.add_paragraph('Course Structure', style='Heading 1')
    Course_competence = competence.runs[0]
    Course_competence.font.size = Pt(16)
    Course_competence.font.name = 'Arial'
    Course_competence.bold = True
    Course_competence.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black


    Semester = doc.add_paragraph('Year 1 – Semester 1 & 2', style='Heading 2')
    Sem = Semester.runs[0]
    Sem.font.size = Pt(12)
    Sem.font.name = 'Arial'
    Sem.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black


    # Create the table with 4 columns (Core Learning Units, Hours, Credits)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Core Learning Units'
    hdr_cells[1].text = 'Hours'
    hdr_cells[2].text = 'Credits'

    # Define the content for Year 1 Semester 1 & 2
    year_1_content = [
        ('ITB111', 'UX Design', 60, 4),
        ('ITB211', 'Statistical Research Methods', 60, 4),
        ('ITB511', 'Programming', 60, 4),
        ('ITB611', 'Network Administration', 30, 2),
        ('ITB811', 'Business Needs Analysis', 30, 2),
        ('ITB411', 'Data Modelling', 30, 2),
        ('ITB221', 'Decision Analysis', 60, 4),
        ('ITB521', 'Data Structures & Algorithms', 60, 4),
        ('ITB621', 'Operating Systems Administration', 60, 4),
        ('ITB731', 'Data Visualisation', 30, 2),
        ('ITB911', 'Applied Cryptography', 30, 2),
        ('ITB421', 'Data Storage Administration', 30, 2),
        ('ITBW21', 'Visual Analytics Project', 30, 2),
        ('', 'General Studies', '', '')
    ]

    # Populate the table for Year 1
    for unit_code, unit_name, hours, credits in year_1_content:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{unit_code} {unit_name}"
        row_cells[1].text = str(hours)
        row_cells[2].text = str(credits)

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')




    # Add "Year 2 – Semester 1 & 2" as a heading
    doc.add_paragraph('Year 2 – Semester 1 & 2', style='Heading 2').runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Create the table for Year 2 – Semester 1 & 2
    table2 = doc.add_table(rows=1, cols=3)
    table2.autofit = True

    # Add the header row for Year 2
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Core Learning Units'
    hdr_cells[1].text = 'Hours'
    hdr_cells[2].text = 'Credits'

    # Define the content for Year 2 Semester 1 & 2
    year_2_content = [
        ('ITB231', 'Supervised Learning', 60, 4),
        ('ITB531', 'Web Application Development', 60, 4),
        ('ITB341', 'Data Wrangling', 30, 2),
        ('ITB641', 'Data Integration & Clustering', 30, 2),
        ('ITB141', 'Data Journalism', 30, 2),
        ('ITB232', 'Unsupervised Learning', 30, 2),
        ('ITB931', 'Data Privacy & Protection', 30, 2),
        ('ITB441', 'Predictive Analytics Project', 30, 2),
        ('ITB251', 'Topic Modelling & Sentiment Analysis', 60, 4),
        ('ITB841', 'Big Data Modelling & Management', 60, 4),
        ('ITB541', 'Programming for Data Science', 60, 4),
        ('ITB721', 'Natural Language Processing', 30, 2),
        ('ITB711', 'Emerging Technology Synthesis', 30, 2),
        ('ITB831', 'Customer Experience Analysis', 30, 2),
        ('ITBW51', 'Text & Social Analytics Project', 30, 2),
        ('', 'General Studies', '', '')
    ]

    # Populate the table for Year 2
    for unit_code, unit_name, hours, credits in year_2_content:
        row_cells = table2.add_row().cells
        row_cells[0].text = f"{unit_code} {unit_name}"
        row_cells[1].text = str(hours)
        row_cells[2].text = str(credits)

    # Optional: Format the table font
    for table in [table, table2]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'


    doc.add_paragraph('Year 3 – Semester 1 & 2', style='Heading 2').runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Create the table for Year 3 – Semester 1 & 2
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Core Learning Units'
    hdr_cells[1].text = 'Hours'
    hdr_cells[2].text = 'Credits'

    # Define the content for Year 3 Semester 1 & 2
    year_3_content = [
        ('IT3301', 'Applied Machine Learning', 60, 4),
        ('IT3381', 'Applied Deep Learning', 30, 2),
        ('IT3382', 'Advanced Data Visualisation', 30, 2),
        ('IT3383', 'Data Processing on Big Data', 30, 2),
        ('IT3384', 'Data Platform Management', 30, 2),
        ('IT3385', 'Machine Learning Operations', 30, 2),
        ('IT3386', 'AI Services in Analytics', 30, 2),
        ('IT3387', 'Marketing Strategy', 30, 2),
        ('IT3331', 'Final Year Project', 480, 12),
        ('IT3336', 'Internship Programme', 480, 12),
        ('IT3333', 'Overseas Internship Programme', 480, 12),
        ('IT3337', 'Final Year Project (24-week)', 960, 24),
        ('IT3339', 'Internship Programme (24-week)', 960, 24),
        ('IT3338', 'Overseas Internship Programme (24-week)', 960, 24),
        ('', 'Prescribed Elective', 30, 2),
        ('', 'General Studies', '', '')
    ]

    # Populate the table for Year 3
    for unit_code, unit_name, hours, credits in year_3_content:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{unit_code} {unit_name}"
        row_cells[1].text = str(hours)
        row_cells[2].text = str(credits)

    # Add "Electives" section heading
    electives_heading = doc.add_paragraph('Electives', style='Heading 1')
    electives_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Add "Prescribed Elective" as a subheading
    prescribed_elective_heading = doc.add_paragraph('Prescribed Elective', style='Heading 2')
    prescribed_elective_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Create the table for Electives
    table2 = doc.add_table(rows=1, cols=3)
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Core Learning Units'
    hdr_cells2[1].text = 'Hours'
    hdr_cells2[2].text = 'Credits'

    ##elective are not found in the top file
    elective_content = [
        ('IT3388', 'Big Data Management Project', 30, 2),
        ('IT3389', 'Applied AI Project', 30, 2)
    ]

    for unit_code, unit_name, hours, credits in elective_content:
        row_cells = table2.add_row().cells
        row_cells[0].text = f"{unit_code} {unit_name}"
        row_cells[1].text = str(hours)
        row_cells[2].text = str(credits)
    for table in [table, table2]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'

    ##Manual Mapping Tables for CMS ##
    custom = doc.add_paragraph('Mapping Tables for Communication Skills & Mathematics Topics')
    custom_run = custom.runs[0]
    custom_run.font.size = Pt(16)
    custom_run.font.name = 'Arial'
    custom_run.bold = True

    #TODO: Frontend
    ##End mapping tabling for cms##




    doc.add_paragraph('ITB111 UX DESIGN', style='Heading 2')

    # Add course details table (Course, Course Code, Year, Duration, etc.)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Populate course details rows
    details_content = [
        ('Course:', 'Diploma in Applied AI & Analytics', 'Course Code:', 'ITDPEFA'),
        ('Year:', '1', 'Duration / Credits:', '60 Hrs / 4'),
        ('Pre/Co-requisite:', 'Nil', '', ''),
        ('Async Lecture (AL):', '15', 'Practical (P):', '42'),
        ('Tutorial (T):', '0', 'eLearning (E):', '3'),
    ]

    for i, (label1, value1, label2, value2) in enumerate(details_content):
        row_cells = table.rows[i].cells
        row_cells[0].text = label1
        row_cells[1].text = value1
        row_cells[2].text = label2
        row_cells[3].text = value2

    # Merge cells for the last row (Pre/Co-requisite row)
    row_cells = table.rows[2].cells
    row_cells[1].merge(row_cells[3])

    # Merge cells for the next duration row
    row_cells = table.rows[5].cells
    row_cells[0].merge(row_cells[1])
    row_cells[2].merge(row_cells[3])

    # Add Synopsis section
    doc.add_paragraph('Synopsis', style='Heading 2')
    doc.add_paragraph(
        "User Experience (UX) provides a positive experience through any form of human-computer interaction. "
        "Through this unit, learners will develop competencies in designing and creating intuitive user interfaces..."
    )

    # Add Learning Outcomes section
    doc.add_paragraph('Learning Outcomes', style='Heading 2')
    doc.add_paragraph(
        'At the end of this unit, learners will be able to:\n'
        '1. Use storyboarding approach to conceptualise and communicate...\n'
        '2. Design intuitive and accessible user interfaces...\n'
        '3. Undertake a technical lead role to drive the team...'
    )

    # Add Topics section
    doc.add_paragraph('Topics', style='Heading 2')
    doc.add_paragraph(
        '1. Importance of UX/UI in web design\n'
        '2. Principles of visual design\n'
        '3. Web accessibility guidelines\n'
        '4. Web client-server architecture\n'
        '5. HTML Document, Elements, Attributes\n'
        '6. CSS Rules and Selectors\n'
        '7. Interactive Web Development\n'
        '8. JavaScript and Control Structures\n'
        '9. JavaScript functions and event handlers\n'
        '10. JavaScript HTML DOM\n'
        '11. Design principles and heuristics\n'
        '12. How to conduct usability testing\n'
        '13. Communication Skills – Oral presentation...'
    )

    # Add Key Tasks section
    doc.add_paragraph('Key Tasks', style='Heading 2')
    doc.add_paragraph(
        '1. Design and build client-based, user-centered web pages using HTML and CSS\n'
        '2. Design and build user-centered web forms using interaction design principles\n'
        '3. Creating responsive layout for web pages\n'
        '4. Communicate the proposed user interactions and experience to key stakeholders\n'
        '5. Build interaactive web pages using Javascript\n'
        '6. Communicate the proposed wrebsite to the key stakeholders'
    )

    doc.add_paragraph('Assessments', style='Heading 2')

    # Create table for assessments
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # Fill in assessment data
    assessments = [
        ('Assignment', '35%'),
        ('Practical', '20%'),
        ('Project', '35%'),
        ('Presentation', '10%'),
    ]

    # Populate the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Assessment'
    hdr_cells[1].text = '%'

    for i, (assessment, percentage) in enumerate(assessments, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = assessment
        row_cells[1].text = percentage

    # Add heading for Texts & References
    doc.add_paragraph('\nTexts & References', style='Heading 2')

    # Add references as a numbered list
    references = [
        'Lean UX: Designing Great Products with Agile Teams, 2016, Jeff Gothelf, Josh Seiden',
        'A Project Guide to UX Design, 2009, Russ Unger, Carolyn Chandler',
        'HTML Comprehensive Concepts and Techniques, Shelly, Woods, Dorin, 5th Edition, CT, B',
        'Teach Yourself Visually HTML5, 2011, Mike Wooldridge, JW, B',
        'Brilliant HTML5 & CSS3, 2011, Hill, Josh, PH, B',
    ]

    for i, reference in enumerate(references, 1):
        doc.add_paragraph(f'{i}. {reference}', style='List Number')



    # Save the document to an in-memory file
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Go back to the beginning of the BytesIO object
    return doc_io

def extract_text_between_sections(file_path, start_section, end_section):
    # Open the Word document
    doc = Document(file_path)
    
    # Variables to control extraction
    extracting = False
    extracted_text = []
    
    for paragraph in doc.paragraphs:
        # Start extracting after the start section
        if start_section in paragraph.text:
            extracting = True
            continue
        
        # Stop extracting when reaching the end section
        if extracting and end_section in paragraph.text:
            break
        
        # Append text if we're in the extraction range
        if extracting:
            extracted_text.append(paragraph.text)
    
    # Join extracted paragraphs into a single string
    return "\n".join(extracted_text)
