import os
import base64
from flask import Flask, render_template, send_file, request
from werkzeug.utils import secure_filename
import zipfile              
from io import BytesIO        
from PIL import Image      
from openpyxl import load_workbook 
from docx import Document                   
from docx.shared import Pt, Inches, RGBColor  
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from collections import defaultdict
import re



app = Flask(__name__)
ALLOWED_EXTENSIONS = {'docx', 'xlsx'}


def extract_images_as_base64(file_stream, file_extension):
    """Extract images from an uploaded DOCX or XLSX file as base64-encoded strings."""
    images_base64 = []

    if file_extension == 'docx':
        with zipfile.ZipFile(file_stream, 'r') as z:
            for file_name in z.namelist():
                if file_name.startswith('word/media/'):
                    with z.open(file_name) as source_file:
                        image_data = source_file.read()
                        image_b64 = base64.b64encode(image_data).decode('utf-8')
                        images_base64.append((file_name, image_b64))
    elif file_extension == 'xlsx':
        workbook = load_workbook(file_stream)
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for image in worksheet._images:
                image_data = image._data()
                image_b64 = base64.b64encode(image_data).decode('utf-8')
                images_base64.append((image.anchor._from, image_b64))

    return images_base64
# Helper function to check allowed files
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def extract_tables_from_docx(doc_path):
    """Extract tables from a DOCX file as python-docx Table objects."""
    doc = Document(doc_path)
    tables = []
    for table in doc.tables:
        tables.append(table)
    return tables
def get_base64_image_from_file(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode('utf-8')

@app.route("/")
@app.route("/index")
def index():
    return render_template("index.html")

@app.route("/upload", methods=['POST'])
def upload_file():
    # Check if both files are in the request
    if 'docx_file' not in request.files or 'xlsx_file' not in request.files:
        return "Both DOCX and XLSX files are required", 400

    docx_file = request.files['docx_file']
    xlsx_file = request.files['xlsx_file']

    # Validate both files
    if docx_file.filename == '' or xlsx_file.filename == '':
        return "No selected file(s)", 400

    if allowed_file(docx_file.filename) and allowed_file(xlsx_file.filename):
        try:
            topics = request.form.getlist('topics[]')
            learning_units = request.form.getlist('learning_units[]')
            docx_extension = docx_file.filename.rsplit('.', 1)[1].lower()
            xlsx_extension = xlsx_file.filename.rsplit('.', 1)[1].lower()

            # Ensure correct file types
            if docx_extension != 'docx' or xlsx_extension != 'xlsx':
                return "Incorrect file types. Please upload a DOCX file and an XLSX file.", 400

            # Save the DOCX file temporarily for extraction
            docx_filename = secure_filename(docx_file.filename)
            temp_docx_path = os.path.join("/tmp", docx_filename)
            docx_file.save(temp_docx_path)

            # Extract images from the uploaded DOCX file
            images_base64 = extract_images_as_base64(temp_docx_path, docx_extension)

            # Extract tables from the uploaded DOCX file
            extracted_tables = extract_tables_from_docx(temp_docx_path)

            # Get the base64 image of the logo from /logo.png
            logo_base64 = get_base64_image_from_file("public/static/images/logo.png")

            # Create document with logo image and XLSX data
            doc_io = create_document(images_base64, logo_base64, xlsx_file, temp_docx_path, topics, learning_units, extracted_tables)

            # Remove the temporary DOCX file after processing
            os.remove(temp_docx_path)

            return send_file(doc_io, as_attachment=True, download_name='output.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            return f"An error occurred during processing: {str(e)}", 500

    return "File types not allowed. Please upload valid DOCX and XLSX files.", 400


def extract_images_as_base64(file_stream, file_extension):
    """Extract images from an uploaded DOCX or XLSX file as base64-encoded strings."""
    images_base64 = []

    if file_extension == 'docx':
        with zipfile.ZipFile(file_stream, 'r') as z:
            for file_name in z.namelist():
                if file_name.startswith('word/media/'):
                    with z.open(file_name) as source_file:
                        image_data = source_file.read()
                        image_b64 = base64.b64encode(image_data).decode('utf-8')
                        images_base64.append((file_name, image_b64))
    elif file_extension == 'xlsx':
        workbook = load_workbook(file_stream)
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for image in worksheet._images:
                image_data = image._data()
                image_b64 = base64.b64encode(image_data).decode('utf-8')
                images_base64.append((image.anchor._from, image_b64))

    return images_base64


def create_document(images_base64, base64_img_first, file, word_doc_path, topics, learning_units, tables):
    """Create a DOCX document with a provided base64 image first and extracted base64 images."""
    doc = Document()
    image_data = base64.b64decode(base64_img_first)
    image_stream = BytesIO(image_data)
        # Open image with PIL to determine size
    with Image.open(image_stream) as img:
        max_width = 6.0  # Max width in inches
        width, height = img.size
        aspect_ratio = width / height

        if width > height:
            adjusted_width = min(max_width, width / 96)  # Convert pixels to inches (assuming 96 dpi)
            adjusted_height = adjusted_width / aspect_ratio
        else:
            adjusted_height = min(max_width, height / 96)
            adjusted_width = adjusted_height * aspect_ratio

        # Add the image to the document
        doc.add_picture(image_stream, width=Inches(adjusted_width), height=Inches(adjusted_height))
        def calculate_column_widths(table):
            """
            Calculate proportional widths for each column based on the text length.
            """
            # Dictionary to store total text length per column
            column_text_length = defaultdict(int)

            # Calculate the total text length for each column
            for row in table.rows:
                for col_idx, cell in enumerate(row.cells):
                    cell_text_length = len(cell.text)
                    column_text_length[col_idx] += cell_text_length

            # Calculate total text length across all columns
            total_text_length = sum(column_text_length.values())

            # Calculate proportional width for each column
            column_widths = {}
            for col_idx, length in column_text_length.items():
                proportion = length / total_text_length if total_text_length else 1 / len(column_text_length)
                column_widths[col_idx] = Inches(6 * proportion)  # Assuming total width of 6 inches

            return column_widths

        def apply_column_widths(table, column_widths):
            """
            Apply calculated widths to each column in the table.
            """
            for row in table.rows:
                for col_idx, cell in enumerate(row.cells):
                    if col_idx in column_widths:
                        cell.width = column_widths[col_idx]  # Set column width proportionally

        # Insert tables into the document after introductory text
        doc.add_paragraph("Below are the tables extracted from the original document:", style='Heading 2')


    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')


    file_extension_school = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_school == 'xlsx':
        workbook_school = load_workbook(file)
        cover_page_school = workbook_school['Cover Page']
        diploma_text = cover_page_school['C6'].value.upper()
        if diploma_text.startswith("IT"):
            it_paragraph = doc.add_paragraph("Diploma in Information Technology")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("BM"):
            it_paragraph = doc.add_paragraph("Diploma in Business Management")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("DM"):
            it_paragraph = doc.add_paragraph("Diploma in Design & Media")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("HS"):
            it_paragraph = doc.add_paragraph("Diploma in Health & Social Sciences")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        elif diploma_text.startswith("EG"):
            it_paragraph = doc.add_paragraph("Diploma in Engineering")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'
        else:
            it_paragraph = doc.add_paragraph("Diploma in NYP")
            it_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            it_run = it_paragraph.runs[0]
            it_run.font.size = Pt(24)
            it_run.font.name = 'Arial'


    course = doc.add_paragraph('Course Syllabi for')
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = course.runs[0]
    run1.font.size = Pt(24)
    run1.font.name = 'Arial'


    file_extension = file.filename.rsplit('.', 1)[1].lower()
    if file_extension == 'xlsx':
        workbook = load_workbook(file)
        cover_page = workbook['Cover Page']
        diploma_text = cover_page['C7'].value.upper()

    diploma = doc.add_paragraph(diploma_text)
    diploma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = diploma.runs[0]
    run2.font.size = Pt(24)
    run2.font.name = 'Arial'
    run2.bold = True


    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    file_extension_year = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_year == 'xlsx':
        workbook_year = load_workbook(file)
        cover_page_year = workbook_year['Cover Page']
        diploma_text_year = cover_page_year['C12'].value.upper()
        year = doc.add_paragraph(diploma_text_year)
        year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = diploma.runs[0]
        run3.font.size = Pt(18)
        run3.font.name = 'Arial'


    doc.add_paragraph('')
    doc.add_paragraph('')


    revised = doc.add_paragraph('Date Revised: Apr  2023 v1.0')
    run4 = revised.runs[0]
    run4.font.size = Pt(11)
    run4.font.name = 'Arial'


    title1 = doc.add_paragraph(diploma_text)
    run5 = title1.runs[0]
    run5.font.size = Pt(16)
    run5.font.name = 'Arial'
    run5.bold = True
        

    title2 = doc.add_paragraph('Course Aims')
    run8 = title2.runs[0]
    run8.font.size = Pt(16)
    run8.font.name = 'Arial'
    run8.bold = True
    doc.add_paragraph('')
    doc.add_paragraph('')




    start_section = "Course Aims"
    end_section = "Course Learning Outcomes"
    text_between_sections = extract_text_between_sections(word_doc_path, start_section, end_section)
    
    course_paragraph = doc.add_paragraph(text_between_sections)
    course_run = course_paragraph.runs[0]
    course_run.font.size = Pt(12)
    course_run.font.name = 'Arial'


    title3 = doc.add_paragraph('Course Learning Outcomes')
    run10 = title3.runs[0]
    run10.font.size = Pt(16)
    run10.font.name = 'Arial'
    run10.bold = True

    course_paragraph_outcome = doc.add_paragraph("The competencies of a graduate are synthesized into 9 Course Competencies (CCs) as listed below.")
    course_run_outcome = course_paragraph_outcome.runs[0]
    course_run_outcome.font.size = Pt(12)
    course_run_outcome.font.name = 'Arial'


    for image_filename, image_b64 in images_base64:

        title2 = doc.add_paragraph('Course Competency Map')
        run8 = title2.runs[0]
        run8.font.size = Pt(16)
        run8.font.name = 'Arial'
        run8.bold = True
        # Decode base64 image and open with PIL
        image_data = base64.b64decode(image_b64)
        image_stream = BytesIO(image_data)

        # Open image with PIL to determine size
        with Image.open(image_stream) as img:
            max_width = 6.0  # Max width in inches
            width, height = img.size
            aspect_ratio = width / height

            if width > height:
                adjusted_width = min(max_width, width / 96)  # Convert pixels to inches (assuming 96 dpi)
                adjusted_height = adjusted_width / aspect_ratio
            else:
                adjusted_height = min(max_width, height / 96)
                adjusted_width = adjusted_height * aspect_ratio

            doc.add_picture(image_stream, width=Inches(adjusted_width), height=Inches(adjusted_height))


    clo_competency = doc.add_paragraph('Course Learning Outcomes')
    run_clo = clo_competency.runs[0]
    run_clo.font.size = Pt(16)
    run_clo.font.name = 'Arial'
    run_clo.bold = True
    table1 = tables[0]

    column_width = calculate_column_widths(table1)

    new_table1 = doc.add_table(rows=len(table1.rows), cols=len(table1.columns))
    new_table1.style = 'Table Grid'

    for i, row in enumerate(table1.rows):
        new_table1.cell(i, 0).text = f"CLO{i + 1}"
        
        for j, cell in enumerate(row.cells):
            if j > 0:  # Skip the first column s it's already set to "CLO1", "CLO2", etc.
                new_table1.cell(i, j).text = cell.text
            
            for paragraph in new_table1.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(new_table1, column_width)
    doc.add_paragraph('')



    title_competency = doc.add_paragraph('Competency Canvases')
    run_course = title_competency.runs[0]
    run_course.font.size = Pt(16)
    run_course.font.name = 'Arial'
    run_course.bold = True
    doc.add_paragraph('')



    para5 = doc.add_paragraph('The learning outcomes of the diploma are to educate and train students who, by the time of successful completion of course, will be able to:')
    run13 = para5.runs[0]
    run13.font.size = Pt(12)
    run13.font.name = 'Arial'


    competency_table = tables[2]
    column_width = calculate_column_widths(competency_table)
    competency_new_table = doc.add_table(rows=len(competency_table.rows), cols=len(competency_table.columns))
    competency_new_table.style = 'Table Grid'

    for i, row in enumerate(competency_table.rows):
        # Start inserting "CLO" labels from the second row (index 1)
        if i > 0:
            competency_new_table.cell(i, 0).text = f"CC{i}"  # Start counting from "CLO1" on the second row
        
        for j, cell in enumerate(row.cells):
            if j > 0:  # Skip the first column as it's already set for CLO
                competency_new_table.cell(i, j).text = cell.text
            
            # Set font size and font name for each cell
            for paragraph in competency_new_table.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(competency_new_table, column_width)
    doc.add_paragraph('')



    competence = doc.add_paragraph('Course Structure', style='Heading 1')
    Course_competence = competence.runs[0]
    Course_competence.font.size = Pt(16)
    Course_competence.font.name = 'Arial'
    Course_competence.bold = True
    Course_competence.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black


    Semester = doc.add_paragraph('Year 1 – Semester 1 & 2', style='Heading 2')
    Sem = Semester.runs[0]
    Sem.font.size = Pt(12)
    Sem.font.name = 'Arial'
    Sem.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black


    table_year1 = tables[10]
    column_width = calculate_column_widths(table_year1)
    new_table_year1 = doc.add_table(rows=len(table_year1.rows), cols=len(table_year1.columns))
    new_table_year1.style = 'Table Grid'
    for i, row in enumerate(table_year1.rows):
        for j, cell in enumerate(row.cells):
            new_table_year1.cell(i, j).text = cell.text
            for paragraph in new_table_year1.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(new_table_year1, column_width)
    doc.add_paragraph('')

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')




    # Add "Year 2 – Semester 1 & 2" as a heading
    doc.add_paragraph('Year 2 – Semester 1 & 2', style='Heading 2').runs[0].font.color.rgb = RGBColor(0, 0, 0)

    #
    table_year2 = tables[11]
    column_width = calculate_column_widths(table_year2)
    new_table_year2 = doc.add_table(rows=len(table_year2.rows), cols=len(table_year2.columns))
    new_table_year2.style = 'Table Grid'
    for i, row in enumerate(table_year2.rows):
        for j, cell in enumerate(row.cells):
            new_table_year2.cell(i, j).text = cell.text
            for paragraph in new_table_year2.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(new_table_year2, column_width)
    doc.add_paragraph('')

    doc.add_paragraph('Year 3 – Semester 1 & 2', style='Heading 2').runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table_year3 = tables[12]
    column_width = calculate_column_widths(table_year3)
    new_table_year3 = doc.add_table(rows=len(table_year3.rows), cols=len(table_year3.columns))
    new_table_year3.style = 'Table Grid'
    for i, row in enumerate(table_year3.rows):
        for j, cell in enumerate(row.cells):
            new_table_year3.cell(i, j).text = cell.text
            for paragraph in new_table_year3.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(new_table_year3, column_width)
    doc.add_paragraph('')

    ##Manual Mapping Tables for CMS ##
    custom = doc.add_paragraph('Mapping Tables for Communication Skills & Mathematics Topics')
    custom_run = custom.runs[0]
    custom_run.font.size = Pt(16)
    custom_run.font.name = 'Arial'
    custom_run.bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

            # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Topic'
    hdr_cells[1].text = 'Learning Unit'

            # Populate table rows with topics and learning units
    for topic, unit in zip(topics, learning_units):
        row_cells = table.add_row().cells
        row_cells[0].text = topic
        row_cells[1].text = unit




    doc.add_paragraph('ITB111 UX DESIGN', style='Heading 2')

    table_year2 = tables[15]
    column_width = calculate_column_widths(table_year2)
    new_table_year2 = doc.add_table(rows=len(table_year2.rows), cols=len(table_year2.columns))
    new_table_year2.style = 'Table Grid'
    for i, row in enumerate(table_year2.rows):
        for j, cell in enumerate(row.cells):
            new_table_year2.cell(i, j).text = cell.text
            for paragraph in new_table_year2.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'

    apply_column_widths(new_table_year2, column_width)
    doc.add_paragraph('')

    # Add Synopsis section
    doc.add_paragraph('Synopsis', style='Heading 2')
    file_extension_Synopsis = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_Synopsis == 'xlsx':
        workbook_Synopsis = load_workbook(file)
        cover_page_Synopsis = workbook_Synopsis['Syllabus Summary']
        diploma_text_Synopsis = cover_page_Synopsis['B12'].value.upper()

    diploma = doc.add_paragraph(diploma_text_Synopsis)
    doc.add_paragraph(diploma_text_Synopsis)

    # Add Learning Outcomes section
    doc.add_paragraph('Learning Outcomes', style='Heading 2')
    doc.add_paragraph('At the end of this unit, learners will be able to:')
    file_extension_outcome = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_outcome == 'xlsx':
        workbook_outcome = load_workbook(file)
        cover_page_outcome = workbook_outcome['Syllabus Summary']
        diploma_text_outcome1 = cover_page_outcome['C13'].value.upper()
        doc.add_paragraph(diploma_text_outcome1)
        diploma_text_outcome2 = cover_page_outcome['C14'].value.upper()
        doc.add_paragraph(diploma_text_outcome2)
        diploma_text_outcome3 = cover_page_outcome['C15'].value.upper()
        doc.add_paragraph(diploma_text_outcome3)

    # Add Topics section
    doc.add_paragraph('Topics', style='Heading 2')
    file_extension_topics = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_topics == 'xlsx':
        workbook_topics = load_workbook(file)
        cover_page_topics = workbook_topics['Syllabus']
        # Start iterating from row 5 in column 'C'
        for cell in cover_page_topics['C5:C{}'.format(cover_page_topics.max_row)]:
            for cell_obj in cell:  # cell_obj is each cell within the specified range
                if cell_obj.value:  # Check if the cell is not empty
                    diploma_text_topics = cell_obj.value.upper()
                    doc.add_paragraph(diploma_text_topics)

    # Add Key Tasks section
    doc.add_paragraph('Key Tasks', style='Heading 2')
    file_extension_tasks = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_tasks == 'xlsx':
        workbook_tasks = load_workbook(file)
        cover_page_tasks = workbook_tasks['Syllabus']
        # Start iterating from row 5 in column 'C'
        for cell in cover_page_tasks['C5:C{}'.format(cover_page_tasks.max_row)]:
            for cell_obj in cell:  # cell_obj is each cell within the specified range
                if cell_obj.value:  # Check if the cell is not empty
                    diploma_text_tasks = cell_obj.value.upper()
                    doc.add_paragraph(diploma_text_tasks)

    doc.add_paragraph('Assessments', style='Heading 2')

    # Assuming `file` is an uploaded XLSX file
    file_extension_grading = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_grading == 'xlsx':
        workbook_grading = load_workbook(file)
        cover_page_grading = workbook_grading['Syllabus']
        
        # Extract the value from cell M5 and convert it to uppercase
        assessments = cover_page_grading['M5'].value.upper()
        
        # Remove "E.G." if present
        assessments = assessments.replace("E.G.", "").strip()

        # Split the assessments text into separate tasks based on "ASSN TASK"
        tasks = [task.strip() for task in assessments.split("ASSN TASK") if task.strip()]
        
        # Create a table with three columns: Assessment, Details, Percentage
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Set header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Assessment'
        hdr_cells[1].text = 'Details'
        hdr_cells[2].text = 'Percentage'

        # Iterate over each task and add it as a new row in the table
        for task in tasks:
            # Add "Assn Task" back to the start of each task for context
            task_text = "Assn Task " + task
            
            # Use regular expression to separate percentage if it exists
            match = re.search(r'(\d+%)', task_text)
            if match:
                # Extract the percentage and the remaining text
                percentage = match.group(1)
                details = task_text.replace(percentage, "").strip()
            else:
                # No percentage found, set percentage to empty and keep full text in details
                percentage = ""
                details = task_text

            # Add a new row to the table for each task
            row_cells = table.add_row().cells
            row_cells[0].text = 'Assessment Details'
            row_cells[1].text = details
            row_cells[2].text = percentage
    doc.add_paragraph('\nTexts & References', style='Heading 2')

    # Add references as a numbered list

    file_extension_references = file.filename.rsplit('.', 1)[1].lower()
    if file_extension_references == 'xlsx':
        workbook_references = load_workbook(file)
        cover_page_references = workbook_references['Syllabus Summary']
        diploma_text_references = cover_page_references['B16'].value.upper()

    diploma = doc.add_paragraph(diploma_text_references)


    # Save the document to an in-memory file
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Go back to the beginning of the BytesIO object
    return doc_io

def extract_text_between_sections(file_path, start_section, end_section):
    # Open the Word document
    doc = Document(file_path)
    
    # Variables to control extraction
    extracting = False
    extracted_text = []
    
    for paragraph in doc.paragraphs:
        # Start extracting after the start section
        if start_section in paragraph.text:
            extracting = True
            continue
        
        # Stop extracting when reaching the end section
        if extracting and end_section in paragraph.text:
            break
        
        # Append text if we're in the extraction range
        if extracting:
            extracted_text.append(paragraph.text)
    
    # Join extracted paragraphs into a single string
    return "\n".join(extracted_text)

if __name__ == '__main__':
    app.run(debug=True)
