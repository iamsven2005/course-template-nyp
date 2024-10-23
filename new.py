from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add a title
title = doc.add_heading('Document Creation Example', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a paragraph with bold and italic text
paragraph = doc.add_paragraph('This is a sample document created using the python-docx library.')
run = paragraph.runs[0]
run.bold = True
run.italic = True

# Add a heading
doc.add_heading('Section 1: Introduction', level=2)

# Add a bulleted list
list_paragraph = doc.add_paragraph()
list_paragraph.add_run('Bullet 1').bold = True
list_paragraph.add_run(' - This is the first bullet point.')
list_paragraph.add_run('\n')
list_paragraph.add_run('Bullet 2').bold = True
list_paragraph.add_run(' - This is the second bullet point.')

# Add a table
doc.add_heading('Section 2: Data', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.autofit = False
table.allow_autofit = False
for row in table.rows:
    for cell in row.cells:
        cell.width = Pt(100)
table.cell(0, 0).text = 'Name'
table.cell(0, 1).text = 'Age'
table.cell(0, 2).text = 'City'
for i, data in enumerate([('Alice', '25', 'New York'), ('Bob', '30', 'San Francisco'), ('Charlie', '22', 'Los Angeles')], start=1):
    table.cell(i, 0).text = data[0]
    table.cell(i, 1).text = data[1]
    table.cell(i, 2).text = data[2]

# Add an image
doc.add_heading('Section 3: Image', level=2)
doc.add_paragraph('Here is an image:')
doc.add_picture('image.jpeg', width=Pt(300))

# Save the document
doc.save('example_document.docx')